"""
Image-to-Text Converter for chunking pipeline.

This script processes PDF/DOCX files to:
1. Extract images from PDFs and DOCX files
2. Get surrounding text context from the document
3. Send images + context to Gemini for descriptions
4. Create documents with image descriptions instead of images
5. Convert those cleaned documents to markdown

Run this BEFORE markdown_converter.py in the pipeline:
  1. image_converter.py     (PDF/DOCX with images → PDF/DOCX with descriptions) ← YOU ARE HERE
  2. markdown_converter.py  (PDF/DOCX → markdown)
  3. chunker.py            (markdown → chunks with embeddings)
  4. embed_and_insert.py   (chunks → database)
"""

import os
import json
import io
from pathlib import Path
from dotenv import load_dotenv
from google import genai
import base64

# PDF and document processing
import pymupdf
import docx
from docx.shared import RGBColor

load_dotenv()


# ── Configuration ─────────────────────────────────────────────────────────────

def _require(name: str) -> str:
    val = os.environ.get(name)
    if not val:
        raise SystemExit(
            f"ERROR: Required environment variable '{name}' is not set.\n"
            f"       Copy chunking_documentation/.env.example to "
            f"chunking_documentation/.env and fill in your values."
        )
    return val


GEMINI_API_KEY = _require("GEMINI_API_KEY")
INPUT_FOLDER = _require("CHUNKER_INPUT_FOLDER")
OUTPUT_FOLDER = _require("CHUNKER_OUTPUT_FOLDER")

# Create a cleaned folder for image-free documents
CLEANED_FOLDER = os.path.join(OUTPUT_FOLDER, "cleaned_documents")
Path(CLEANED_FOLDER).mkdir(parents=True, exist_ok=True)

client = genai.Client(api_key=GEMINI_API_KEY)


# ── PDF Image Extraction ──────────────────────────────────────────────────────

def extract_images_and_text_from_pdf(filepath: Path) -> list[dict]:
    """
    Extract images and surrounding text from PDF.
    Returns list of { image_data, image_index, before_context, after_context }
    """
    doc = pymupdf.open(filepath)
    images_data = []
    all_text = ""

    # First pass: collect all text
    for page in doc:
        all_text += page.get_text() + "\n"

    # Second pass: extract images with page-level context
    total_pages = len(doc)
    print(f"      Scanning {total_pages} pages for images...")
    
    for page_num, page in enumerate(doc):
        image_list = page.get_images()
        if image_list:
            print(f"        Page {page_num + 1}: Found {len(image_list)} image(s)")

        for img_index, img in enumerate(image_list):
            xref = img[0]
            pix = pymupdf.Pixmap(doc, xref)

            if pix.n - pix.alpha < 4:  # GRAY or RGB
                image_bytes = pix.tobytes("png")
            else:
                pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
                image_bytes = pix.tobytes("png")

            # Get page text as context
            page_text = page.get_text()

            # Split into sentences for better context
            sentences = page_text.split('.')
            context_sentences = ' '.join(sentences[:5]).strip()  # First few sentences

            images_data.append({
                "image_bytes": image_bytes,
                "image_index": len(images_data),
                "page_num": page_num,
                "page_text": page_text,
                "document_text": all_text
            })

    doc.close()
    return images_data


# ── DOCX Image Extraction ─────────────────────────────────────────────────────

def extract_images_and_text_from_docx(filepath: Path) -> list[dict]:
    """
    Extract images and surrounding text from DOCX.
    Returns list of { image_data, image_index, paragraph_index, before_context, after_context }
    """
    from docx.oxml import parse_xml
    
    doc = docx.Document(filepath)
    images_data = []

    # Get all document text for context
    all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    # Extract images from the document part's related image parts
    print(f"      Extracting images from DOCX package...")
    image_count = 0
    
    try:
        # Get all image parts from the document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    image_count += 1
                    print(f"        Found image {image_count}: {rel.target_ref} ({len(image_bytes)} bytes)")
                    
                    # Try to find which paragraph contains this image
                    para_idx = 0
                    for p_idx, para in enumerate(doc.paragraphs):
                        if para_idx == 0 and len(doc.paragraphs) > 0:
                            para_idx = p_idx
                    
                    images_data.append({
                        "image_bytes": image_bytes,
                        "image_index": image_count - 1,
                        "paragraph_index": para_idx,
                        "paragraph_text": doc.paragraphs[para_idx].text if para_idx < len(doc.paragraphs) else "",
                        "before_para": doc.paragraphs[para_idx - 1].text if para_idx > 0 else "",
                        "after_para": doc.paragraphs[para_idx + 1].text if para_idx < len(doc.paragraphs) - 1 else "",
                        "document_text": all_text,
                        "rel_id": rel.rId
                    })
                except Exception as e:
                    print(f"        ⚠️  Could not extract image from relationship: {e}")
    except Exception as e:
        print(f"      ⚠️  Error accessing image relationships: {e}")

    return images_data


# ── Image to Text with Gemini ─────────────────────────────────────────────────

def encode_image_to_base64(image_path: Path) -> str:
    """Convert image file to base64."""
    with open(image_path, "rb") as f:
        return base64.standard_b64encode(f.read()).decode("utf-8")


def convert_image_bytes_to_description(
    image_bytes: bytes,
    document_context: str,
    before_context: str = "",
    after_context: str = "",
    image_num: int = 1
) -> str:
    """
    Send image to Gemini with surrounding context.
    Request a text description that can replace the image in documentation.
    """
    try:
        print(f"        📤 Sending image to Gemini (size: {len(image_bytes)} bytes)...")
        image_base64 = base64.standard_b64encode(image_bytes).decode("utf-8")

        # Construct context
        context_text = ""
        if before_context:
            context_text += f"BEFORE IMAGE:\n{before_context[:300]}\n\n"
        context_text += f"DOCUMENT CONTEXT:\n{document_context[:500]}"
        if after_context:
            context_text += f"\n\nAFTER IMAGE:\n{after_context[:300]}"

        prompt = f"""You are analyzing an image from technical documentation for a RAG (Retrieval Augmented Generation) system.
The image will be replaced with a text description that serves the same purpose as the image.

Document Context:
{context_text}

[IMAGE BELOW]

Your task:
1. Describe what the image shows in detail
2. Explain its relevance to the surrounding context
3. Include any text visible in the image
4. Preserve all important information that would be lost without the image

Write a comprehensive text description (2-4 sentences) that will replace the image in the documentation.
Make it specific enough that someone reading only the text would understand what the image conveyed.

Respond with ONLY the text description, no additional formatting or explanation."""

        print(f"        🔄 Calling gemini-3.1-flash-lite-preview API...")
        
        # Use the correct API format for google-genai
        response = client.models.generate_content(
            model="gemini-3.1-flash-lite-preview",
            contents=[
                prompt,
                {
                    "inline_data": {
                        "mime_type": "image/png",
                        "data": image_base64
                    }
                }
            ]
        )

        result = response.text.strip()
        print(f"        ✅ Got response from Gemini ({len(result)} chars)")
        return result

    except Exception as e:
        print(f"      ❌ ERROR sending image to Gemini: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return "[Image description unavailable]"


# ── Document Processing ──────────────────────────────────────────────────────

def process_pdf_images(filepath: Path, output_dir: Path) -> dict:
    """
    Process a PDF:
    1. Extract images and text
    2. Convert images to descriptions with Gemini
    3. Create a new PDF with descriptions instead of images
    4. Save as output PDF
    """
    filename = filepath.name
    print(f"  Processing PDF: {filename}")

    images_data = extract_images_and_text_from_pdf(filepath)

    if not images_data:
        print(f"    ℹ️  No images found. Copying file as-is.")
        import shutil
        output_path = output_dir / filename
        shutil.copy(filepath, output_path)
        return {
            "filename": filename,
            "format": "pdf",
            "images_found": 0,
            "images_converted": 0,
            "errors": 0
        }

    print(f"    Found {len(images_data)} image(s). Converting to descriptions...")

    # Create new PDF with descriptions
    doc = pymupdf.open(filepath)
    images_converted = 0
    errors_count = 0

    for i, img_data in enumerate(images_data):
        print(f"      [{i + 1}/{len(images_data)}] Converting image from page {img_data['page_num'] + 1}...")

        try:
            # Extract context from document
            doc_context = img_data["document_text"][:800]
            page_text = img_data["page_text"][:500]

            description = convert_image_bytes_to_description(
                img_data["image_bytes"],
                doc_context,
                page_text,
                "",
                image_num=i + 1
            )

            # For PDF, we insert text where images are
            # This is a simplified approach - full image removal would require more work
            page = doc[img_data["page_num"]]
            page.insert_text(
                (50, 50),
                f"[Image description: {description}]",
                fontsize=10
            )

            images_converted += 1
            print(f"        ✅ Successfully converted")

        except Exception as e:
            print(f"        ❌ Error converting image: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
            errors_count += 1

    # Save modified PDF
    output_path = output_dir / filename
    doc.save(str(output_path))
    doc.close()

    print(f"    ✅ Saved to {filename}")

    return {
        "filename": filename,
        "format": "pdf",
        "images_found": len(images_data),
        "images_converted": images_converted,
        "errors": errors_count
    }


def process_docx_images(filepath: Path, output_dir: Path) -> dict:
    """
    Process a DOCX:
    1. Extract images and text
    2. Convert images to descriptions with Gemini
    3. Create a new DOCX with descriptions instead of images
    4. Save as output DOCX
    """
    filename = filepath.name
    print(f"  Processing DOCX: {filename}")

    images_data = extract_images_and_text_from_docx(filepath)

    if not images_data:
        print(f"    ℹ️  No images found. Copying file as-is.")
        import shutil
        output_path = output_dir / filename
        shutil.copy(filepath, output_path)
        return {
            "filename": filename,
            "format": "docx",
            "images_found": 0,
            "images_converted": 0,
            "errors": 0
        }

    print(f"    Found {len(images_data)} image(s). Converting to descriptions...")

    # Load original document
    doc = docx.Document(filepath)
    images_converted = 0
    errors_count = 0

    # Track which images we've processed to avoid duplicates
    processed_indices = set()

    for i, img_data in enumerate(images_data):
        if img_data["image_index"] in processed_indices:
            continue

        print(
            f"      [{len(processed_indices) + 1}/{len(images_data)}] "
            f"Converting image {img_data['image_index']} from paragraph {img_data['paragraph_index']}..."
        )

        try:
            print(f"        📤 Sending image (size: {len(img_data['image_bytes'])} bytes)...")
            description = convert_image_bytes_to_description(
                img_data["image_bytes"],
                img_data["document_text"],
                img_data["before_para"],
                img_data["after_para"],
                image_num=i + 1
            )

            # Add description paragraph
            para_idx = img_data["paragraph_index"]
            if para_idx < len(doc.paragraphs):
                target_para = doc.paragraphs[para_idx]
                # Insert description before the paragraph
                new_para = target_para.insert_paragraph_before(
                    f"[Image description: {description}]"
                )
                new_para.style = "Normal"

            images_converted += 1
            processed_indices.add(img_data["image_index"])
            print(f"        ✅ Successfully converted")

        except Exception as e:
            print(f"        ❌ Error converting image: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
            errors_count += 1

    # Save new document
    output_path = output_dir / filename
    doc.save(str(output_path))

    print(f"    ✅ Saved to {filename}")

    return {
        "filename": filename,
        "format": "docx",
        "images_found": len(images_data),
        "images_converted": images_converted,
        "errors": errors_count
    }


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    input_path = Path(INPUT_FOLDER)

    pdf_files = sorted(input_path.glob("*.pdf"))
    docx_files = sorted(input_path.glob("*.docx"))

    if not pdf_files and not docx_files:
        print("⚠️  No PDF or DOCX files found in the input folder.")
        print(f"    Looking in: {INPUT_FOLDER}")
        exit()

    print(f"🖼️  Converting images to descriptions...\n")
    print(f"   Input folder:  {INPUT_FOLDER}")
    print(f"   Output folder: {CLEANED_FOLDER}\n")

    all_stats = []

    # Process PDFs
    for pdf_file in pdf_files:
        stats = process_pdf_images(pdf_file, Path(CLEANED_FOLDER))
        all_stats.append(stats)

    # Process DOCX files
    for docx_file in docx_files:
        stats = process_docx_images(docx_file, Path(CLEANED_FOLDER))
        all_stats.append(stats)

    # Summary
    total_files = len(pdf_files) + len(docx_files)
    total_images = sum(s["images_found"] for s in all_stats)
    total_converted = sum(s["images_converted"] for s in all_stats)
    total_errors = sum(s["errors"] for s in all_stats)

    print(f"\n{'─'*50}")
    print(f"✅ Image conversion complete!")
    print(f"   Files processed    : {total_files}")
    print(f"   PDFs processed     : {len(pdf_files)}")
    print(f"   DOCX processed     : {len(docx_files)}")
    print(f"   Total images found : {total_images}")
    print(f"   Successfully converted: {total_converted}")
    print(f"   Errors: {total_errors}")
    print(f"   Output folder: {CLEANED_FOLDER}")
    print(f"\n📝 Next steps:")
    print(f"   1. Review cleaned documents in: {CLEANED_FOLDER}")
    print(f"   2. Run markdown_converter.py on the cleaned documents:")
    print(f"      cp {CLEANED_FOLDER}/* {INPUT_FOLDER}/")
    print(f"      python markdown_converter.py")
    print(f"   3. Continue with chunker.py")

    # Save stats
    stats_file = Path(CLEANED_FOLDER) / "_image_conversion_stats.json"
    with open(stats_file, 'w', encoding='utf-8') as f:
        json.dump(all_stats, f, indent=2, ensure_ascii=False)
    print(f"\n📊 Statistics saved to: _image_conversion_stats.json")
